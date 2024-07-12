import os
from typing import List, Tuple, Dict
from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def add_table_to_doc(
    doc: DocxDocument,
    table_data: List[List[str]],
    colored_cells_for_table: Tuple[int, int, int],
) -> None:
    """Добавляет таблицу в документ.

    :param doc: Объект документа.
    :param table_data: Данные таблицы.
    :param colored_cells_for_table: Координаты ячеек для закрашивания.
    """
    header = table_data[0]
    rows = table_data[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    # Заполняем заголовок таблицы
    fill_table_header(table, header)

    # Заполняем строки таблицы
    fill_table_rows(table, rows)

    # Убираем границы у первой колонки (кроме правой)
    remove_first_column_borders(table)

    # Закраска определенной ячейки
    if colored_cells_for_table:
        _, idx_row, idx_col = colored_cells_for_table
        if len(table.rows) > idx_row and len(table.rows[1].cells) > idx_col:
            color_table_cell(table, idx_row, idx_col)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    doc.add_paragraph()


def fill_table_header(table: Table, header: List[str]) -> None:
    """Заполняет заголовок таблицы.

    :param table: Объект таблицы.
    :param header: Данные заголовка.
    """
    hdr_cells = table.rows[0].cells
    for idx, header_text in enumerate(header):
        hdr_cells[idx].text = header_text
        # Выравнивание заголовка
        if idx == 0:
            hdr_cells[idx].paragraphs[
                0
            ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            hdr_cells[idx].paragraphs[
                0
            ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def fill_table_rows(table: Table, rows: List[List[str]]) -> None:
    """Заполняет строки таблицы.

    :param table: Объект таблицы.
    :param rows: Данные строк.
    """
    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells  # Получаем объект ячеек строки
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(
                row_cells
            ):  # Проверяем, что индекс в пределах допустимых значений
                row_cells[col_idx].text = cell_text
            else:
                row_cells[-1].text = (
                    cell_text  # Если вышли за пределы, записываем в последнюю ячейку
                )
            # Выравнивание текста в ячейках
            if col_idx == 0:
                row_cells[col_idx].paragraphs[
                    0
                ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                row_cells[col_idx].paragraphs[
                    0
                ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def remove_first_column_borders(table: Table) -> None:
    """Убирает границы у первой колонки (кроме правой).

    :param table: Объект таблицы.
    """
    for row in table.rows:
        cell = row.cells[0]
        tcPr = cell._element.get_or_add_tcPr()
        tcPr.append(
            parse_xml(
                (
                    "<w:tcBorders {}>"
                    '<w:top w:val="nil"/>'
                    '<w:left w:val="nil"/>'
                    '<w:bottom w:val="nil"/>'
                    '<w:right w:val="single"/>'
                    "</w:tcBorders>"
                ).format(nsdecls("w"))
            )
        )


def color_table_cell(table: Table, row_idx: int, col_idx: int) -> None:
    """Закрашивает ячейку таблицы желтым цветом.

    :param table: Объект таблицы.
    :param row_idx: Индекс строки.
    :param col_idx: Индекс столбца.
    """
    cell_to_color = table.rows[row_idx].cells[col_idx]
    tcPr = cell_to_color._element.get_or_add_tcPr()
    tcPr.append(parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls("w"))))


def write_to_docx(
    filename: str,
    tables_data: Dict[str, Tuple[List[List[str]], Tuple[int, int, int]]],
) -> None:
    """Записывает данные из таблиц в файл формата .docx.

    :param filename: Имя файла .docx для записи.
    :param tables_data: Словарь с данными таблиц в формате {имя_файла:
        (tables, colored_cells)}.
    """
    docx_filepath = os.path.join(os.getcwd(), filename)

    doc = Document()

    for filename, (tables, colored_cells) in tables_data.items():
        data_from = doc.add_heading(f"Данные из файла {filename}", level=1)
        data_from.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, (table_data, colored_cells_for_table) in enumerate(tables):
            heading = doc.add_heading(f"Таблица {idx + 1}", level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            add_table_to_doc(doc, table_data, colored_cells_for_table)

        doc.add_page_break()

    doc.save(docx_filepath)
